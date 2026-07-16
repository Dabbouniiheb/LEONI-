#!/usr/bin/env python3
"""Génère les formats HTML et DOCX du rapport, sans toucher à l'application."""

from __future__ import annotations

import html
import re
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
MD_PATH = BASE_DIR / "rapport_stage_complet.md"
HTML_PATH = BASE_DIR / "rapport_stage_complet.html"
DOCX_PATH = BASE_DIR / "rapport_stage_complet.docx"

PLACEHOLDER_RE = re.compile(r"\[(?:INFORMATION À COMPLÉTER|CAPTURE À INSÉRER|ÉLÉMENT À CONFIRMER)\s*:[^\]]+\]")
SOURCE_COMMENT_RE = re.compile(r"^\s*<!--\s*Sources? (?:projet|code)\s*:.*?-->\s*$", re.I | re.S)


def slugify(value: str) -> str:
    value = value.lower()
    replacements = str.maketrans(
        {
            "à": "a",
            "â": "a",
            "ä": "a",
            "ç": "c",
            "é": "e",
            "è": "e",
            "ê": "e",
            "ë": "e",
            "î": "i",
            "ï": "i",
            "ô": "o",
            "ö": "o",
            "ù": "u",
            "û": "u",
            "ü": "u",
            "ÿ": "y",
            "œ": "oe",
        }
    )
    value = value.translate(replacements)
    value = re.sub(r"[^a-z0-9]+", "-", value).strip("-")
    return value or "section"


def parse_table_row(line: str) -> list[str]:
    body = line.strip().strip("|")
    return [cell.strip() for cell in body.split("|")]


def is_table_separator(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def is_special(line: str, next_line: str | None = None) -> bool:
    stripped = line.strip()
    if not stripped:
        return True
    if stripped == "<!-- PAGE_BREAK -->":
        return True
    if SOURCE_COMMENT_RE.match(stripped):
        return True
    if stripped.startswith("#") or stripped.startswith("```") or stripped.startswith(">"):
        return True
    if re.match(r"^\s*[-*+]\s+", line) or re.match(r"^\s*\d+[.)]\s+", line):
        return True
    if "|" in line and next_line and is_table_separator(next_line):
        return True
    return False


def inline_html(text: str) -> str:
    tokens: dict[str, str] = {}

    def stash(rendered: str) -> str:
        key = f"@@TOKEN{len(tokens)}@@"
        tokens[key] = rendered
        return key

    text = re.sub(
        r"<((?:https?://)[^>]+)>",
        lambda m: stash(f'<a href="{html.escape(m.group(1), quote=True)}">{html.escape(m.group(1))}</a>'),
        text,
    )
    text = re.sub(
        r"\[([^\]]+)\]\((https?://[^)]+)\)",
        lambda m: stash(f'<a href="{html.escape(m.group(2), quote=True)}">{html.escape(m.group(1))}</a>'),
        text,
    )
    text = re.sub(
        r"`([^`]+)`",
        lambda m: stash(f"<code>{html.escape(m.group(1))}</code>"),
        text,
    )
    text = PLACEHOLDER_RE.sub(
        lambda m: stash(f'<span class="placeholder">{html.escape(m.group(0))}</span>'), text
    )
    text = html.escape(text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<em>\1</em>", text)
    for key, rendered in tokens.items():
        text = text.replace(key, rendered)
    return text


def collect_headings(lines: list[str]) -> list[tuple[int, str, str]]:
    seen: dict[str, int] = {}
    result: list[tuple[int, str, str]] = []
    for line in lines:
        match = re.match(r"^(#{1,3})\s+(.+?)\s*$", line)
        if not match:
            continue
        level = len(match.group(1))
        title = re.sub(r"[`*_]", "", match.group(2))
        base = slugify(title)
        seen[base] = seen.get(base, 0) + 1
        anchor = base if seen[base] == 1 else f"{base}-{seen[base]}"
        result.append((level, title, anchor))
    return result


def render_html(lines: list[str]) -> str:
    headings = collect_headings(lines)
    heading_iter = iter(headings)
    current_heading = next(heading_iter, None)
    body: list[str] = []
    i = 0
    in_cover = True

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if SOURCE_COMMENT_RE.match(stripped):
            i += 1
            continue
        if stripped == "<!-- PAGE_BREAK -->":
            body.append('<div class="page-break"></div>')
            in_cover = False
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            body.append(f'<pre data-language="{html.escape(language)}"><code>{html.escape(chr(10).join(code_lines))}</code></pre>')
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.+?)\s*$", raw)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            if current_heading is None:
                anchor = slugify(title)
            else:
                anchor = current_heading[2]
                current_heading = next(heading_iter, None)
            css = ' class="cover-title"' if in_cover and level <= 2 else ""
            body.append(f'<h{level} id="{anchor}"{css}>{inline_html(title)}</h{level}>')
            if re.sub(r"[`*_]", "", title).strip().lower() == "table des matières":
                toc = ['<nav class="toc"><ol>']
                for h_level, h_title, h_anchor in headings:
                    if h_title.lower() == "table des matières" or h_level > 3:
                        continue
                    toc.append(
                        f'<li class="toc-level-{h_level}"><a href="#{h_anchor}">{inline_html(h_title)}</a></li>'
                    )
                toc.append("</ol></nav>")
                body.extend(toc)
            i += 1
            continue
        if "|" in raw and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header_cells = parse_table_row(raw)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip() and "|" in lines[i]:
                rows.append(parse_table_row(lines[i]))
                i += 1
            body.append('<div class="table-wrap"><table><thead><tr>')
            body.extend(f"<th>{inline_html(cell)}</th>" for cell in header_cells)
            body.append("</tr></thead><tbody>")
            for row in rows:
                row = row + [""] * max(0, len(header_cells) - len(row))
                body.append("<tr>")
                body.extend(f"<td>{inline_html(cell)}</td>" for cell in row[: len(header_cells)])
                body.append("</tr>")
            body.append("</tbody></table></div>")
            continue
        unordered = re.match(r"^\s*[-*+]\s+(.+)$", raw)
        ordered = re.match(r"^\s*\d+[.)]\s+(.+)$", raw)
        if unordered or ordered:
            tag = "ul" if unordered else "ol"
            body.append(f"<{tag}>")
            while i < len(lines):
                match = re.match(r"^\s*[-*+]\s+(.+)$", lines[i]) if tag == "ul" else re.match(r"^\s*\d+[.)]\s+(.+)$", lines[i])
                if not match:
                    break
                body.append(f"<li>{inline_html(match.group(1))}</li>")
                i += 1
            body.append(f"</{tag}>")
            continue
        if stripped.startswith(">"):
            quote: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip()[1:].strip())
                i += 1
            body.append(f"<blockquote>{inline_html(' '.join(quote))}</blockquote>")
            continue

        paragraph = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            nxt2 = lines[i + 1] if i + 1 < len(lines) else None
            if is_special(nxt, nxt2):
                break
            paragraph.append(nxt.strip())
            i += 1
        css = ' class="cover-line"' if in_cover else ""
        body.append(f"<p{css}>{inline_html(' '.join(paragraph))}</p>")

    css = """
@page { size: A4; margin: 2.2cm 2.4cm 2.2cm 2.5cm; }
:root { --ink:#171717; --muted:#555; --line:#9a9a9a; --shade:#f1f1f1; --accent:#2f3a45; }
* { box-sizing:border-box; }
html { scroll-behavior:smooth; }
body { max-width: 21cm; margin:0 auto; padding:1.2cm; color:var(--ink); background:white; font-family:"Times New Roman", Times, serif; font-size:12pt; line-height:1.5; text-align:justify; }
h1,h2,h3,h4,h5,h6 { color:#111; text-align:left; page-break-after:avoid; break-after:avoid; }
h1 { font-size:18pt; margin:1.2em 0 .65em; border-bottom:1px solid var(--line); padding-bottom:.18em; }
h2 { font-size:14pt; margin:1.15em 0 .45em; }
h3 { font-size:12.5pt; margin:1em 0 .35em; }
p { margin:.35em 0 .75em; orphans:3; widows:3; }
.cover-title { text-align:center; border:0; margin-top:1.1cm; }
.cover-line { text-align:center; margin:.45em 0; }
.page-break { break-before:page; page-break-before:always; height:0; }
.placeholder { background:#fff1a8; border:1px solid #c8ad43; padding:0 .15em; font-weight:600; }
code { font-family:"Courier New", monospace; font-size:.88em; background:#f3f3f3; padding:.08em .22em; }
pre { white-space:pre-wrap; background:#f3f3f3; border-left:3px solid #777; padding:.7em; font-size:9pt; line-height:1.3; text-align:left; break-inside:avoid; }
blockquote { border-left:3px solid #777; margin:.8em 0; padding:.25em .9em; color:#333; }
ul,ol { margin:.35em 0 .8em 1.4em; padding-left:.5em; }
li { margin:.15em 0; }
.table-wrap { overflow-x:auto; margin:.7em 0 1em; }
table { width:100%; border-collapse:collapse; font-size:8.6pt; line-height:1.25; text-align:left; break-inside:auto; }
thead { display:table-header-group; }
tr { break-inside:avoid; page-break-inside:avoid; }
th,td { border:1px solid #777; padding:.28em .38em; vertical-align:top; }
th { background:#dedede; font-weight:700; }
.toc { text-align:left; border:1px solid #aaa; padding:.7em 1em; background:#fafafa; }
.toc ol { list-style:none; margin:0; padding:0; }
.toc-level-2 { margin-left:1.2em; }
.toc-level-3 { margin-left:2.4em; font-size:10.5pt; }
a { color:#1e3e5c; text-decoration:none; }
@media print { body { max-width:none; margin:0; padding:0; } a { color:#111; } .toc { border:0; } }
"""
    return "<!doctype html>\n<html lang=\"fr\"><head><meta charset=\"utf-8\"><meta name=\"viewport\" content=\"width=device-width,initial-scale=1\"><title>Rapport de stage — Planning Home Office</title><style>" + css + "</style></head><body>\n" + "\n".join(body) + "\n</body></html>\n"


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, name: str = "Times New Roman", size: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1E3E5C")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_TOKEN_RE = re.compile(
    r"(\[(?:INFORMATION À COMPLÉTER|CAPTURE À INSÉRER|ÉLÉMENT À CONFIRMER)\s*:[^\]]+\]|\*\*[^*]+\*\*|`[^`]+`|<https?://[^>]+>|\[[^\]]+\]\(https?://[^)]+\))"
)


def add_inline(paragraph, text: str, default_size: float | None = None) -> None:
    pos = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=default_size)
        token = match.group(0)
        if PLACEHOLDER_RE.fullmatch(token):
            run = paragraph.add_run(token)
            set_run_font(run, size=default_size)
            run.bold = True
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "FFF1A8")
            run._r.get_or_add_rPr().append(shading)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size)
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "Courier New", default_size or 9.5)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F2F2F2")
            run._r.get_or_add_rPr().append(shading)
        elif token.startswith("<http"):
            url = token[1:-1]
            add_hyperlink(paragraph, url, url)
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=default_size)


def add_toc_field(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Table des matières actualisable : clic droit > Mettre à jour le champ."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    r_pr.extend([fonts, size])
    r.append(r_pr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.3)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)
    section.different_first_page_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Title": (22, True, WD_ALIGN_PARAGRAPH.CENTER, 18, 12),
        "Subtitle": (15, False, WD_ALIGN_PARAGRAPH.CENTER, 8, 8),
        "Heading 1": (18, True, WD_ALIGN_PARAGRAPH.LEFT, 16, 8),
        "Heading 2": (14, True, WD_ALIGN_PARAGRAPH.LEFT, 12, 5),
        "Heading 3": (12.5, True, WD_ALIGN_PARAGRAPH.LEFT, 9, 4),
        "Heading 4": (12, True, WD_ALIGN_PARAGRAPH.LEFT, 8, 3),
    }
    for style_name, (size, bold, align, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(20, 20, 20)
        style.paragraph_format.alignment = align
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in [s.name for s in styles]:
        style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        style.font.size = Pt(9)
        style.paragraph_format.left_indent = Cm(0.5)
        style.paragraph_format.right_indent = Cm(0.3)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.05

    for sec in document.sections:
        header = sec.header
        hp = header.paragraphs[0]
        hp.text = "Rapport de stage — Application de gestion du planning Home Office"
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hp.runs:
            set_run_font(run, size=9)
            run.font.color.rgb = RGBColor(85, 85, 85)
        footer = sec.footer
        fp = footer.paragraphs[0]
        add_page_number(fp)


def render_docx(lines: list[str]) -> Document:
    document = Document()
    configure_document(document)
    document.core_properties.title = "Rapport de stage — Gestion du planning Home Office"
    document.core_properties.subject = "Rétrodocumentation code-only de l'application leoni-planing"
    document.core_properties.author = "[INFORMATION À COMPLÉTER : nom de l'étudiant]"
    document.core_properties.comments = "Généré depuis rapport_stage_complet.md"

    i = 0
    in_cover = True
    first_heading = True
    pending_page_break = False
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if SOURCE_COMMENT_RE.match(stripped):
            i += 1
            continue
        if stripped == "<!-- PAGE_BREAK -->":
            pending_page_break = True
            in_cover = False
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            p = document.add_paragraph(style="Code Block")
            p.paragraph_format.keep_together = True
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F2F2F2")
            p_pr.append(shd)
            run = p.add_run("\n".join(code_lines))
            set_run_font(run, "Courier New", 9)
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.+?)\s*$", raw)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            if first_heading:
                p = document.add_paragraph(style="Title")
                first_heading = False
            else:
                style_level = min(level, 4)
                p = document.add_paragraph(style=f"Heading {style_level}")
            if pending_page_break:
                p.paragraph_format.page_break_before = True
                pending_page_break = False
            add_inline(p, title)
            if in_cover:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if level == 1:
                    p.paragraph_format.space_before = Pt(34)
                elif level == 2:
                    p.paragraph_format.space_before = Pt(18)
            if re.sub(r"[`*_]", "", title).strip().lower() == "table des matières":
                add_toc_field(document)
            i += 1
            continue
        if "|" in raw and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = parse_table_row(raw)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip() and "|" in lines[i]:
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = document.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            font_size = 8.8 if len(headers) <= 4 else 7.6 if len(headers) >= 7 else 8.1
            for idx, value in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_shading(cell, "D9D9D9")
                set_cell_margins(cell)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_inline(p, value, font_size)
                for run in p.runs:
                    run.bold = True
            set_repeat_table_header(table.rows[0])
            for row_values in rows:
                row_values = row_values + [""] * max(0, len(headers) - len(row_values))
                cells = table.add_row().cells
                for idx, value in enumerate(row_values[: len(headers)]):
                    cell = cells[idx]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    set_cell_margins(cell)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    add_inline(p, value, font_size)
            document.add_paragraph().paragraph_format.space_after = Pt(0)
            continue
        unordered = re.match(r"^\s*[-*+]\s+(.+)$", raw)
        ordered = re.match(r"^\s*(\d+)[.)]\s+(.+)$", raw)
        if unordered or ordered:
            p = document.add_paragraph(style="List Bullet" if unordered else "Normal")
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.35)
            p.paragraph_format.space_after = Pt(2)
            if unordered:
                add_inline(p, unordered.group(1))
            else:
                add_inline(p, f"{ordered.group(1)}. {ordered.group(2)}")
            i += 1
            continue
        if stripped.startswith(">"):
            quote_parts: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_parts.append(lines[i].strip()[1:].strip())
                i += 1
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.right_indent = Cm(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, " ".join(quote_parts))
            continue

        paragraph = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            nxt2 = lines[i + 1] if i + 1 < len(lines) else None
            if is_special(nxt, nxt2):
                break
            paragraph.append(nxt.strip())
            i += 1
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if in_cover else WD_ALIGN_PARAGRAPH.JUSTIFY
        if in_cover:
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        add_inline(p, " ".join(paragraph))

    return document


def main() -> int:
    if not MD_PATH.exists():
        print(f"Fichier source introuvable : {MD_PATH}", file=sys.stderr)
        return 1
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()
    HTML_PATH.write_text(render_html(lines), encoding="utf-8")
    document = render_docx(lines)
    document.save(DOCX_PATH)
    print(f"HTML : {HTML_PATH}")
    print(f"DOCX : {DOCX_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
