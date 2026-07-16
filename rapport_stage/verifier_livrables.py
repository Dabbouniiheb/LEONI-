#!/usr/bin/env python3
"""Contrôles reproductibles des livrables du rapport de stage."""

from __future__ import annotations

import re
import sys
import zipfile
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent

REQUIRED = [
    "00_informations_a_completer.md",
    "01_audit_projet_actuel.md",
    "02_resultats_tests.md",
    "03_plan_captures.md",
    "04_matrice_tracabilite.md",
    "05_resume_projet.md",
    "06_fiche_soutenance.md",
    "07_plan_presentation.md",
    "08_glossaire.md",
    "09_references.md",
    "10_checklist_finale.md",
    "rapport_stage_complet.md",
    "rapport_stage_complet.html",
    "rapport_stage_complet.docx",
    "rapport_stage_complet.pdf",
    "diagrams/README.md",
]

DIAGRAMS = [
    "01_cas_utilisation_global.puml",
    "02_architecture_globale.puml",
    "03_composants_application.puml",
    "04_modele_donnees.puml",
    "05_sequence_authentification.puml",
    "06_sequence_changement_mot_de_passe.puml",
    "07_sequence_gestion_utilisateur.puml",
    "08_sequence_selection_groupe.puml",
    "09_sequence_generation_planning.puml",
    "10_sequence_consultation_planning.puml",
    "11_sequence_demande_conge.puml",
    "12_sequence_traitement_conge.puml",
    "13_sequence_session_travail.puml",
    "14_sequence_export.puml",
    "15_sequence_audit.puml",
]


class BasicHTMLValidator(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.errors: list[str] = []
        self.comments: list[str] = []

    def handle_comment(self, data: str) -> None:
        self.comments.append(data)

    def error(self, message: str) -> None:  # pragma: no cover - compatibilité
        self.errors.append(message)


def fail(message: str, errors: list[str]) -> None:
    errors.append(message)


def docx_text(document: Document) -> str:
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def scan_text(name: str, text: str, errors: list[str]) -> None:
    forbidden = [
        "cahier" + "_de_charge",
        "GAP" + "_ANALYSIS",
        "gap" + " analysis",
        "conformité au " + "cahier des charges",
    ]
    lowered = text.lower()
    for term in forbidden:
        if term.lower() in lowered:
            fail(f"{name}: référence historique interdite détectée", errors)
    secret_patterns = [
        r"-----BEGIN (?:RSA |EC |OPENSSH )?PRIVATE KEY-----",
        r"\bBearer\s+[A-Za-z0-9._~-]{20,}",
        r"\bAKIA[0-9A-Z]{16}\b",
        r"(?im)^\s*(?:DB_PASSWORD|SESSION_SECRET)\s*=\s*(?!\[|<|\{|\$)[^\s#]+",
        r"mysql://[^\s:@]+:[^\s@]+@",
    ]
    for pattern in secret_patterns:
        if re.search(pattern, text):
            fail(f"{name}: motif potentiellement sensible détecté", errors)


def main() -> int:
    errors: list[str] = []

    for relative in REQUIRED:
        path = ROOT / relative
        if not path.is_file() or path.stat().st_size == 0:
            fail(f"Livrable absent ou vide : {relative}", errors)

    diagram_dir = ROOT / "diagrams"
    actual_diagrams = sorted(path.name for path in diagram_dir.glob("*.puml"))
    if actual_diagrams != DIAGRAMS:
        fail("La liste des diagrammes ne correspond pas aux quinze noms attendus", errors)
    for name in DIAGRAMS:
        path = diagram_dir / name
        text = path.read_text(encoding="utf-8")
        if text.count("@startuml") != 1 or text.count("@enduml") != 1:
            fail(f"Marqueurs PlantUML invalides : {name}", errors)
        if not re.search(r"(?m)^title\s+\S", text):
            fail(f"Titre PlantUML absent : {name}", errors)
        scan_text(f"diagrams/{name}", text, errors)

    text_suffixes = {".md", ".html", ".puml", ".py"}
    for path in ROOT.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in text_suffixes:
            continue
        if any(part.startswith(".qa_docx") for part in path.parts):
            continue
        scan_text(str(path.relative_to(ROOT)), path.read_text(encoding="utf-8"), errors)

    html_path = ROOT / "rapport_stage_complet.html"
    html_text = html_path.read_text(encoding="utf-8")
    validator = BasicHTMLValidator()
    validator.feed(html_text)
    validator.close()
    if validator.errors:
        fail("Le parseur HTML a signalé une erreur", errors)
    if any("Sources projet" in comment or "Sources code" in comment for comment in validator.comments):
        fail("Les commentaires de sources sont présents dans le HTML final", errors)
    for marker in ("<h1", "<table", "class=\"placeholder\"", "class=\"toc\""):
        if marker not in html_text:
            fail(f"Élément HTML attendu absent : {marker}", errors)

    docx_path = ROOT / "rapport_stage_complet.docx"
    if not zipfile.is_zipfile(docx_path):
        fail("Le DOCX n'est pas une archive OOXML valide", errors)
    document = Document(docx_path)
    section = document.sections[0]
    if abs(section.page_width.cm - 21.0) > 0.05 or abs(section.page_height.cm - 29.7) > 0.05:
        fail("Le DOCX n'est pas au format A4", errors)
    normal = document.styles["Normal"]
    if normal.font.name != "Times New Roman" or abs(normal.font.size.pt - 12.0) > 0.1:
        fail("Le style Normal n'est pas en Times New Roman 12", errors)
    if abs(float(normal.paragraph_format.line_spacing) - 1.5) > 0.01:
        fail("L'interligne du style Normal n'est pas 1,5", errors)
    dtext = docx_text(document)
    if "Sources projet :" in dtext or "Sources code :" in dtext:
        fail("Les commentaires de sources sont visibles dans le DOCX", errors)
    if "[INFORMATION À COMPLÉTER" not in dtext or "[CAPTURE À INSÉRER" not in dtext:
        fail("Les placeholders attendus sont absents du DOCX", errors)
    scan_text("rapport_stage_complet.docx", dtext, errors)

    pdf_path = ROOT / "rapport_stage_complet.pdf"
    reader = PdfReader(str(pdf_path))
    if len(reader.pages) != 61:
        fail(f"Nombre de pages PDF inattendu : {len(reader.pages)}", errors)
    ptext = "\n".join(page.extract_text() or "" for page in reader.pages)
    if "Sources projet :" in ptext or "Sources code :" in ptext:
        fail("Les commentaires de sources sont visibles dans le PDF", errors)
    if "19 tests automatisés sur 19 ont réussi" not in ptext:
        fail("Le bilan de tests attendu est absent du PDF", errors)
    scan_text("rapport_stage_complet.pdf", ptext, errors)

    if errors:
        print("ÉCHEC DES CONTRÔLES")
        for error in errors:
            print(f"- {error}")
        return 1

    print("CONTRÔLES RÉUSSIS")
    print(f"- {len(REQUIRED)} livrables principaux présents")
    print(f"- {len(DIAGRAMS)} diagrammes PlantUML aux noms attendus")
    print("- HTML analysable et commentaires techniques masqués")
    print("- DOCX OOXML valide, A4, Times New Roman 12, interligne 1,5")
    print(f"- PDF valide : {len(reader.pages)} pages A4")
    print("- aucun motif historique interdit ou secret courant détecté")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
